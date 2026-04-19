"""Document loaders for PDF, plain-text, Word (.docx), and web sources."""

from __future__ import annotations

import datetime
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import httpx
import structlog
from bs4 import BeautifulSoup

from rag.exceptions import DocumentLoadError

logger = structlog.get_logger(__name__)

_STRIP_TAGS = {"nav", "footer", "script", "style", "header", "aside"}


@dataclass
class Document:
    """A loaded document before chunking."""

    text: str
    metadata: dict[str, Any]
    source: str


class PDFLoader:
    """Load a PDF file page-by-page using pypdf."""

    def load(self, path: Path) -> list[Document]:
        """Return one Document per page. Raises DocumentLoadError on any failure."""
        try:
            import pypdf  # noqa: PLC0415
        except ImportError as exc:
            raise DocumentLoadError(
                source=str(path),
                message="pypdf is not installed. Run: pip install pypdf",
            ) from exc

        try:
            reader = pypdf.PdfReader(str(path))
            total_pages = len(reader.pages)
            documents: list[Document] = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = text.strip()
                if not text:
                    logger.debug("pdf_page_empty", source=str(path), page=page_num)
                    continue
                documents.append(
                    Document(
                        text=text,
                        metadata={
                            "source": path.name,
                            "page": page_num,
                            "total_pages": total_pages,
                        },
                        source=str(path),
                    )
                )

            logger.info(
                "pdf_loaded",
                source=str(path),
                pages=total_pages,
                non_empty=len(documents),
            )
            return documents

        except Exception as exc:
            raise DocumentLoadError(source=str(path)) from exc


class TextLoader:
    """Load a plain-text file (.txt, .md, or any UTF-8 text) as a single Document."""

    def load(self, path: Path) -> list[Document]:
        """Return a single Document with the file's full text.

        Raises DocumentLoadError on any read failure.
        """
        try:
            text = path.read_text(encoding="utf-8").strip()
            if not text:
                logger.warning("text_file_empty", source=str(path))
                return []
            logger.info("text_loaded", source=str(path), chars=len(text))
            return [
                Document(
                    text=text,
                    metadata={"source": path.name},
                    source=str(path),
                )
            ]
        except Exception as exc:
            raise DocumentLoadError(source=str(path)) from exc


class DocxLoader:
    """Load a Microsoft Word (.docx) file using python-docx.

    Extracts paragraph text in document order. Tables and headers are included
    as paragraphs; images and embedded objects are silently skipped.
    """

    def load(self, path: Path) -> list[Document]:
        """Return a single Document containing all paragraph text.

        Raises DocumentLoadError on any failure, including missing python-docx.
        """
        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:
            raise DocumentLoadError(
                source=str(path),
                message="python-docx is not installed. Run: pip install python-docx",
            ) from exc

        try:
            doc = docx.Document(str(path))

            paragraphs: list[str] = []
            for para in doc.paragraphs:
                stripped = para.text.strip()
                if stripped:
                    paragraphs.append(stripped)

            # Also extract text from table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        stripped = cell.text.strip()
                        if stripped:
                            paragraphs.append(stripped)

            text = "\n\n".join(paragraphs)
            if not text:
                logger.warning("docx_file_empty", source=str(path))
                return []

            logger.info("docx_loaded", source=str(path), paragraphs=len(paragraphs))
            return [
                Document(
                    text=text,
                    metadata={"source": path.name},
                    source=str(path),
                )
            ]
        except DocumentLoadError:
            raise
        except Exception as exc:
            raise DocumentLoadError(source=str(path)) from exc


class WebLoader:
    """Load a web page using httpx and parse its text with BeautifulSoup."""

    def __init__(self, timeout: float = 30.0) -> None:
        self._timeout = timeout

    def load(self, url: str) -> list[Document]:
        """Return a single Document with the page's cleaned text.

        Strips nav, footer, script, and style tags before extracting text.
        Raises DocumentLoadError on network or parse errors.
        """
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (compatible; Stratum-RAG/0.1; +https://github.com/adedaramola/stratum)"
            )
        }
        try:
            response = httpx.get(url, timeout=self._timeout, follow_redirects=True, headers=headers)
            response.raise_for_status()

            soup = BeautifulSoup(response.text, "html.parser")
            for tag in soup.find_all(_STRIP_TAGS):
                tag.decompose()

            text = soup.get_text(separator=" ", strip=True)
            fetched_at = datetime.datetime.now(datetime.UTC).isoformat()

            logger.info("web_loaded", url=url, chars=len(text))
            return [
                Document(
                    text=text,
                    metadata={"source": url, "fetched_at": fetched_at},
                    source=url,
                )
            ]

        except Exception as exc:
            raise DocumentLoadError(source=url) from exc
